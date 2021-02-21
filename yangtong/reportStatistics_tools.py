#!/usr/bin/env python
# -*- coding: utf-8 -*-

import jieba
import jieba.posseg as pseg
import os
import sys
from sklearn import feature_extraction
from sklearn.feature_extraction.text import TfidfTransformer
from sklearn.feature_extraction.text import CountVectorizer
from docx import Document
import re
import pandas as pd
import numpy as np
from sklearn.cluster import KMeans
from mpl_toolkits.mplot3d import Axes3D
import matplotlib.pyplot as plt
from sklearn.manifold import TSNE
import csv
import codecs


def hello():
    print("Hello world")


def getTestDataSource():
    corpus = ["我 来到 北京 清华大学",  # 第一类文本切词后的结果，词之间以空格隔开
            "他 来到 了 网易 杭研 大厦",  # 第二类文本的切词结果
            "小明 硕士 毕业 与 中国 科学院",  # 第三类文本的切词结果
            "我 爱 北京 天安门"]  # 第四类文本的切词结果

    return corpus


# 获取真实的文档语料，不按年份合并，每个文档是一类
def getDataSourceByDocument(file_dir):
    print("-------in function of getDataSourceByDocument------")
    # 1. 读取路径下所有文件名
    for root, dirs, files in os.walk(file_dir):
        files_name = files
    print("files_name = ", files_name)

    # 2. 根据文件名获取文件内容
    files_content = getContentByDocument(files_name, file_dir)

    return files_content


# 计算词频
def computeWordFrequency(corpus):
    # 将文本中的词语转换为词频矩阵
    vectorizer = CountVectorizer()
    # 计算个词语出现的次数
    X = vectorizer.fit_transform(corpus)
    # 获取词袋中所有文本关键词
    words = vectorizer.get_feature_names()
    print(words)
    # 查看词频结果
    X = X.toarray().T
    print(X)
    print("X.shape = ", X.shape)
    print("X's type = ", type(X))
    # X归一化
    x_min, x_max = X.min(0), X.max(0)
    X = (X - x_min) / (x_max - x_min)  # 归一化

    return X, words


# 将变量内容保存至文件
def saveToFile(X, file_path):
    with codecs.open(file_path, "w", "utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(X)
        f.close()


# 2. 根据文件名获取文件内容
def getContentByDocument(files_name, file_dir):
    files_content = []

    # 按年份读取、汇总文件内容
    for file_name in files_name:
        path = file_dir + "\\" + file_name
        # print("path = ", path)
        document = Document(path)
        current_content = ""
        for paragraph in document.paragraphs:
            current_content += paragraph.text
            # print(paragraph.text)
        files_content.append(current_content)

    # print("files_content = ", files_content)
    # print("files_content's length = ", len(files_content))

    files_content = cleanContentByDocument(files_content)

    return files_content


# 接收词频矩阵，训练模型并返回
def trainKMeansModel(words, word_frequency, n_clusters):
    # 创建KMeans 对象
    cluster = KMeans(n_clusters=n_clusters, random_state=0, n_jobs=-1)
    # print("sentence_embeddings = ", sentence_embeddings)
    result = cluster.fit(word_frequency)

    centers = result.cluster_centers_
    # print("centers = ", centers)
    print("centers' length = ", len(centers))

    labels = result.labels_
    labels_set = set(labels)
    print("labels_set = ", labels_set)
    print("labels' length = ", len(labels))
    print("labels = ", result.labels_)

    words_list = getLabeledWords(labels, words)

    tsne = TSNE(n_components=2, init='pca', random_state=501)
    X_tsne = tsne.fit_transform(word_frequency)
    print("Org data dimension is {}.Embedded data dimension is {}".format(word_frequency.shape[-1], X_tsne.shape[-1]))
    x_min, x_max = X_tsne.min(0), X_tsne.max(0)
    X_norm = (X_tsne - x_min) / (x_max - x_min)  # 归一化
    # X_norm = X_tsne
    # print("X_norm = ", X_norm)
    print("X_norm.shape = ", X_norm.shape)
    plt.figure(figsize=(10, 9))
    for i in range(X_norm.shape[0]):
        # print("word = ", words[i])
        plt.text(X_norm[i, 0], X_norm[i, 1], str(labels[i]), weight='light', color=plt.cm.Set1(labels[i]), fontdict={'weight': 'light', 'size': 9})
        # plt.text(X_norm[i, 0], X_norm[i, 1], str(labels[i]), color=plt.cm.Set1(labels[i]), fontdict={'weight': 'bold', 'size': 9})
    plt.xticks([])
    plt.yticks([])
    plt.show()

    return result


# 根据labels和words输出每个label都有哪些单词
def getLabeledWords(labels, words):
    labels_number = len(set(labels))
    print("labels_number = ", labels_number)
    words_list = {}
    for i in range(len(labels)):
        # if i == 500:
        #     break
        label = str(labels[i])
        if label not in words_list.keys():
            words_list[label] = [words[i]]
        else:
            words_list[label].append(words[i])

    print("words_list = ", words_list)

    return words_list


# 根据聚类结果画出图
def pltShow(kmeans_result):
    label_pred = kmeans_result.labels_  # 获取聚类标签

    centroids = kmeans_result.cluster_centers_  # 获取聚类中心

    print(label_pred)
    print(centroids)

    fig = plt.figure()
    ax = Axes3D(fig)
    ax.scatter(data[:, 0], data[:, 1], data[:, 2], c=y, marker='*')

    ax.scatter(centroids[:, 0], centroids[:, 1], centroids[:, 2], marker='>')
    plt.axis([0, 1, 0, 1])
    plt.show()


# 3. 去标点符号 分词 去停用词
def cleanContentByDocument(files_content):
    print("-------in function of cleanContentByDocument------")
    files_content_new = []
    for file_content in files_content:
        # 去标点符号
        content = re.sub("[\s+\.\!\/_,$%^*(+\"\'～]+|[+——！，。？、~@#￥%……&*（）．《 》\]\[；：〔〕】【|]+", "", str(file_content))
        # 分词
        content = list(jieba.cut(content))
        # print("分词后的content = ", content)
        # print("分词后的content's length = ", len(content))

        # 去停用词，同时将content连接为一个字符串
        stop_list = getStopList()
        final_content = ""
        for word in content:
            if word not in stop_list:
                final_content += word.strip()
                final_content += " "

        files_content_new.append(final_content)

    print("去标点符号之后files_content = ", files_content_new)
    return files_content_new


# 获取真实的文档语料
def getDataSource(file_dir):
    print("-------in function of getDataSource------")

    files_name_16, files_name_17, files_name_18, files_name_19, files_name_20 = getAllFilesNameByYear(file_dir)

    files_content_16, files_content_17, files_content_18, files_content_19, files_content_20 = \
        getContentByFilesName(file_dir, files_name_16, files_name_17, files_name_18, files_name_19, files_name_20)

    # 去标点符号、去停用词、jieba分词
    corpus_16, corpus_17, corpus_18, corpus_19, corpus_20 = \
        cleanContent(files_content_16, files_content_17, files_content_18, files_content_19, files_content_20)
    # print("final corpus_16 = ", corpus_16)
    # print("content_16's length = ", len(content_16))

    corpus = [corpus_16, corpus_17, corpus_18, corpus_19, corpus_20]  # corpus的长度为5，分别代表16-20年每年的文档

    return corpus


# 将文本内容去标点符号、分词、去停用词
def cleanContent(files_content_16, files_content_17, files_content_18, files_content_19, files_content_20):
    # print("files_content_16 = ", files_content_16)

    # 去标点符号
    content_16 = re.sub("[\s+\.\!\/_,$%^*(+\"\'～]+|[+——！，。？、~@#￥%……&*（）．《 》\]\[；：〔〕】【|]+", "", str(files_content_16))
    content_17 = re.sub("[\s+\.\!\/_,$%^*(+\"\'～]+|[+——！，。？、~@#￥%……&*（）．《 》\]\[；：〔〕】【|]+", "", str(files_content_17))
    content_18 = re.sub("[\s+\.\!\/_,$%^*(+\"\'～]+|[+——！，。？、~@#￥%……&*（）．《 》\]\[；：〔〕】【|]+", "", str(files_content_18))
    content_19 = re.sub("[\s+\.\!\/_,$%^*(+\"\'～]+|[+——！，。？、~@#￥%……&*（）．《 》\]\[；：〔〕】【|]+", "", str(files_content_19))
    content_20 = re.sub("[\s+\.\!\/_,$%^*(+\"\'～]+|[+——！，。？、~@#￥%……&*（）．《 》\]\[；：〔〕】【|]+", "", str(files_content_20))

    # print("去标点符号后的content_16 = ", content_16)

    # 分词
    content_16 = list(jieba.cut(content_16))
    content_17 = list(jieba.cut(content_17))
    content_18 = list(jieba.cut(content_18))
    content_19 = list(jieba.cut(content_19))
    content_20 = list(jieba.cut(content_20))

    # print("分词后的content_16 = ", content_16)
    # print("分词后的content_16's length = ", len(content_16))
    # print("分词后的content_17 = ", content_17)
    # print("分词后的content_17's length = ", len(content_17))
    # print("分词后的content_18 = ", content_18)
    # print("分词后的content_18's length = ", len(content_18))
    # print("分词后的content_19 = ", content_19)
    # print("分词后的content_19's length = ", len(content_19))
    # print("分词后的content_20 = ", content_20)
    # print("分词后的content_20's length = ", len(content_20))

    # 去停用词，同时将content连接为一个字符串
    stop_list = getStopList()
    final_content_16 = ""
    for word in content_16:
        if word not in stop_list and not word.isdigit():
            final_content_16 += word.strip()
            final_content_16 += " "
    final_content_17 = ""
    for word in content_17:
        if word not in stop_list and not word.isdigit():
            final_content_17 += word.strip()
            final_content_17 += " "
    final_content_18 = ""
    for word in content_18:
        if word not in stop_list and not word.isdigit():
            final_content_18 += word.strip()
            final_content_18 += " "
    final_content_19 = ""
    for word in content_19:
        if word not in stop_list and not word.isdigit():
            final_content_19 += word.strip()
            final_content_19 += " "
    final_content_20 = ""
    for word in content_20:
        if word not in stop_list and not word.isdigit():
            final_content_20 += word.strip()
            final_content_20 += " "

    # print("final_content_16 = ", final_content_16)
    # final_content_16_list = list(final_content_16)
    # print("final_content_16_list = ", final_content_16_list)
    # print("final_content_16_list's length = ", len(final_content_16_list))

    return final_content_16.strip(), final_content_17.strip(), final_content_18.strip(), final_content_19.strip(), final_content_20.strip()


# 获取停用词
def getStopList():
    stoplist = pd.read_csv('stopwords.txt').values
    return stoplist


# 根据文件名，提取文件内容
def getContentByFilesName(file_dir, files_name_16, files_name_17, files_name_18, files_name_19, files_name_20):
    files_content_16 = ""
    files_content_17 = ""
    files_content_18 = ""
    files_content_19 = ""
    files_content_20 = ""

    # 按年份读取、汇总文件内容
    for file_name in files_name_16:
        path = file_dir + "\\" + file_name
        # print("path = ", path)
        document = Document(path)
        for paragraph in document.paragraphs:
            files_content_16 += paragraph.text
            # print(paragraph.text)
    for file_name in files_name_17:
        path = file_dir + "\\" + file_name
        # print("path = ", path)
        document = Document(path)
        for paragraph in document.paragraphs:
            files_content_17 += paragraph.text
            # print(paragraph.text)
    for file_name in files_name_18:
        path = file_dir + "\\" + file_name
        # print("path = ", path)
        document = Document(path)
        for paragraph in document.paragraphs:
            files_content_18 += paragraph.text
            # print(paragraph.text)
    for file_name in files_name_19:
        path = file_dir + "\\" + file_name
        # print("path = ", path)
        document = Document(path)
        for paragraph in document.paragraphs:
            files_content_19 += paragraph.text
            # print(paragraph.text)
    for file_name in files_name_20:
        path = file_dir + "\\" + file_name
        # print("path = ", path)
        document = Document(path)
        for paragraph in document.paragraphs:
            files_content_20 += paragraph.text
            # print(paragraph.text)

    # print(files_content_20)

    return files_content_16, files_content_17, files_content_18, files_content_19, files_content_20


# 根据路径，按年份输出路径下所有文件名
def getAllFilesNameByYear(file_dir):
    # 1. 读取路径下所有文件名，每个年份的文件合成一类文档
    for root, dirs, files in os.walk(file_dir):
        files_name = files
        # print("files_name = ", files_name)
    files_name_16 = []
    files_name_17 = []
    files_name_18 = []
    files_name_19 = []
    files_name_20 = []
    for name in files_name:
        if name.startswith("2016"):
            files_name_16.append(name)
        elif name.startswith("2017"):
            files_name_17.append(name)
        elif name.startswith("2018"):
            files_name_18.append(name)
        elif name.startswith("2019"):
            files_name_19.append(name)
        elif name.startswith("2020"):
            files_name_20.append(name)
    '''
    print("16:", files_name_16)
    print("17:", files_name_17)
    print("18:", files_name_18)
    print("19:", files_name_19)
    print("20:", files_name_20)
    '''

    return files_name_16, files_name_17, files_name_18, files_name_19, files_name_20


# 计算热词，即tf-idf
# corpus格式为字符串列表，以空格隔开每个单词，corpus = ["我 来到 北京 清华大学", "他 来到 了 网易 杭研 大厦"]
def computeHotWords(corpus, limit):
    print("-------in function of computeHotWord------")

    vectorizer = CountVectorizer()  # 该类会将文本中的词语转换为词频矩阵，矩阵元素a[i][j] 表示j词在i类文本下的词频
    transformer = TfidfTransformer()  # 该类会统计每个词语的tf-idf权值
    tfidf = transformer.fit_transform(vectorizer.fit_transform(corpus))  # 第一个fit_transform是计算tf-idf，第二个fit_transform是将文本转为词频矩阵
    word = vectorizer.get_feature_names()  # 获取词袋模型中的所有词语
    print("word's length = ", len(word))
    weights = tfidf.toarray()  # 将tf-idf矩阵抽取出来，元素a[i][j]表示j词在i类文本中的tf-idf权重

    # weights归一化
    # weights_min, weights_max = weights.min(0), weights.max(0)
    # weights = (weights - weights_min) / (weights_max - weights_min)  # 归一化

    results = getHotWords(word, weights, limit)

    '''
    for i in range(len(weight)):  # 打印每类文本的tf-idf词语权重，第一个for遍历所有文本，第二个for便利某一类文本下的词语权重
        print("-------这里输出第", i, "类文本的词语tf-idf权重------")
        for j in range(len(word)):
            print(word[j], weight[i][j])
    '''

    print("-------end of computeHotWord function------")

    return results, word, weights.T


# 根据word和weights取出前50的热词
def getHotWords(word, weights, limit):
    results = []
    for weight in weights:
        results.append(getTop50HotWords(word, weight, limit))

    return results


# 根据word和单个weight计算前50热词，结果以字典形式保存result={word:weight}
def getTop50HotWords(word, weight, limit):
    word_length = len(word)
    weight_length = len(weight)
    if word_length != weight_length:
        print("出问题啦！！！word和weight长度不一致！！！检查一下！！！")
        sys.exit(-1)

    my_dict = dict(zip(word, weight))
    # print("my_dict = ", my_dict)
    my_dict_sorted = sorted(my_dict.items(), key=lambda kv: (kv[1], kv[0]), reverse=True)

    # print("my_dict_sorted = ", my_dict_sorted)
    # print("my_dict_sorted_top50 = ", list(my_dict_sorted)[: limit])

    result = list(my_dict_sorted)[: limit]

    return result


# 计算新词，参考文献
def computeNewAndInheritWords(data_source, limit, words, tfidf_weights):
    print("-------in function of computeNewWord------")

    # 0. 将data_source改为数组形式（当前为字符串形式）
    data = dataSourceSplit(data_source)

    # 1. 计算新颖度、传承度，二维数组
    novelty_weights, inherit_weights = computeNoveltyAndInherit(data, words)

    # 3. 获取tf-idf,tfidf_weights
    # 4. 计算新词系数，三者相乘
    weights = novelty_weights * inherit_weights * tfidf_weights

    results = getHotWords(words, weights, limit)

    print("-------end of computeNewWord function------")

    return results


# 计算新颖度
def computeNoveltyAndInherit(data, words):
    length = len(data)
    words_length = len(words)
    novelty = [[0 for col in range(words_length)] for row in range(length)]
    inherit = [[0 for col in range(words_length)] for row in range(length)]
    print("pre novelty.shape = ", np.array(novelty).shape)
    for i in range(length):
        print("i = ", i)
        if i == 0 or i == (length - 1):
            continue
        pre_data = data[i - 1]
        pos_data = data[i + 1]
        for j in range(words_length):
            if words[j] not in pre_data:
                novelty[i][j] = 1
            else:
                novelty[i][j] = 0.2
            if words[j] not in pos_data:
                inherit[i][j] = 1/3
            else:
                inherit[i][j] = 1

    # print("novelty's length = ", len(novelty))
    # print("novelty's shape = ", np.array(novelty).shape)
    # print("inherit's length = ", len(inherit))
    # print("inherit's shape = ", np.array(inherit).shape)

    return np.array(novelty), np.array(inherit)


# 将字符串改为数组，按空格split，保存在set中
def dataSourceSplit(data_source):
    data_source = data_source
    result = []
    for data in data_source:
        result.append(set(data.strip().split(" ")))

    return result

